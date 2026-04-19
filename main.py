#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Parser for tourist attractions from DOCX files
Extracts data, enriches with AI (OpenRouter), finds photos via Yandex/Google
"""

import os
import re
import json
import time
import requests
import subprocess
import tempfile
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
import hashlib
from datetime import datetime
from urllib.parse import urlparse

# DOCX parsing
from docx import Document

# AI client
from openai import OpenAI

@dataclass
class Attraction:
    """Data class for tourist attraction"""
    name: str
    district: str
    settlement: str = ""
    description_html: str = ""
    photo_url: str = ""
    photo_source: str = ""
    photo_confidence: float = 0.0
    additional_info: Dict = field(default_factory=dict)
    
    def to_dict(self):
        result = {
            'name': self.name,
            'district': self.district,
            'settlement': self.settlement,
            'description_html': self.description_html,
            'photo_url': self.photo_url,
            'photo_source': self.photo_source,
            'photo_confidence': self.photo_confidence,
        }
        if self.additional_info:
            result['additional_info'] = self.additional_info
        return result

class DistrictNormalizer:
    """Normalize district names using provided mapping"""
    
    def __init__(self, districts_json: str = None):
        self.district_map = {}  # normalized_name -> official_name
        self.settlement_map = {}  # district_name -> [settlements]
        
        if districts_json:
            self.load_from_json(districts_json)
    
    def load_from_json(self, json_data):
        """Load districts and settlements from JSON"""
        if isinstance(json_data, str):
            try:
                data = json.loads(json_data)
            except:
                data = []
        else:
            data = json_data
        
        districts = {}
        settlements = {}
        
        for item in data:
            if item.get('model') == 'locations.District':
                districts[item['pk']] = item['fields']['name']
            elif item.get('model') == 'locations.Settlement':
                dist_id = item['fields'].get('district')
                settlement_name = item['fields']['name']
                if dist_id and dist_id in districts:
                    dist_name = districts[dist_id]
                    if dist_name not in settlements:
                        settlements[dist_name] = []
                    settlements[dist_name].append(settlement_name)
        
        # Create normalization map
        for dist_name in districts.values():
            normalized = self._normalize_name(dist_name)
            self.district_map[normalized] = dist_name
        
        self.settlement_map = settlements
    
    def _normalize_name(self, name: str) -> str:
        """Normalize district name for matching"""
        name = name.lower().strip()
        # Remove common suffixes
        for suffix in [' район', 'district', 'р-н']:
            name = name.replace(suffix, '')
        return name.strip()
    
    def normalize(self, folder_name: str) -> str:
        """Normalize folder name to official district name"""
        normalized = self._normalize_name(folder_name)
        
        # Try exact match
        if normalized in self.district_map:
            return self.district_map[normalized]
        
        # Try partial match
        for key, value in self.district_map.items():
            if key in normalized or normalized in key:
                return value
        
        # Return cleaned folder name
        return folder_name.replace(' район', '').strip() + 'ский район'
    
    def find_settlement(self, district: str, text: str) -> str:
        """Find settlement name in text for given district"""
        settlements = self.settlement_map.get(district, [])
        
        for settlement in settlements:
            if settlement.lower() in text.lower():
                return settlement
        
        return ""

class DOCXParser:
    """Parser for DOCX files with tables support"""
    
    @staticmethod
    def _convert_doc_to_docx(file_path: str) -> Optional[str]:
        """Convert legacy .doc file to .docx using LibreOffice"""
        with tempfile.TemporaryDirectory(prefix="doc_convert_") as temp_dir:
            command = [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                temp_dir,
                file_path,
            ]
            result = subprocess.run(command, capture_output=True, text=True, timeout=90)
            if result.returncode != 0:
                return None
            converted = Path(temp_dir) / f"{Path(file_path).stem}.docx"
            if not converted.exists():
                return None
            final_path = Path(tempfile.gettempdir()) / f"{Path(file_path).stem}_{int(time.time() * 1000)}.docx"
            final_path.write_bytes(converted.read_bytes())
            return str(final_path)

    @staticmethod
    def extract_content(file_path: str) -> Tuple[str, List[Dict], str, Optional[str]]:
        """
        Extract text, tables and metadata from DOCX
        Returns: (full_text, tables_data, title, error_message)
        """
        source_file = file_path
        temp_converted_file = None
        try:
            extension = Path(file_path).suffix.lower()
            if extension == ".doc":
                temp_converted_file = DOCXParser._convert_doc_to_docx(file_path)
                if not temp_converted_file:
                    return "", [], "", "Не удалось конвертировать .doc в .docx через soffice"
                source_file = temp_converted_file

            doc = Document(source_file)
            paragraphs = []
            tables_data = []
            title = ""
            
            # Extract title from first paragraph if it looks like a title
            if doc.paragraphs:
                first_para = doc.paragraphs[0].text.strip()
                if len(first_para) < 200 and first_para:
                    title = first_para
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'index': table_idx,
                    'rows': []
                }
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    if row_data:
                        table_data['rows'].append(row_data)
                        # Keep table text in full-text fallback for LLM extraction
                        paragraphs.append(" | ".join(row_data))
                
                if table_data['rows']:
                    tables_data.append(table_data)
            
            full_text = "\n".join(paragraphs)
            return full_text, tables_data, title, None
            
        except Exception as e:
            return "", [], "", str(e)
        finally:
            if temp_converted_file and os.path.exists(temp_converted_file):
                try:
                    os.remove(temp_converted_file)
                except OSError:
                    pass

class ImageSearcher:
    """Search images using API-first strategy with validation"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept-Language': 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
        })
        self.serpapi_key = os.getenv("SERPAPI_API_KEY", "")
        self.google_cse_key = os.getenv("GOOGLE_CSE_API_KEY", "")
        self.google_cse_cx = os.getenv("GOOGLE_CSE_CX", "")
        self.retries = 2

    def _request_json(self, url: str, params: Dict, timeout: int = 12) -> Optional[Dict]:
        for attempt in range(self.retries + 1):
            try:
                response = self.session.get(url, params=params, timeout=timeout)
                if response.status_code != 200:
                    time.sleep(0.5)
                    continue
                if "application/json" not in response.headers.get("Content-Type", ""):
                    time.sleep(0.5)
                    continue
                return response.json()
            except Exception:
                if attempt < self.retries:
                    time.sleep(0.5)
        return None

    def _validate_image_url(self, image_url: str, min_bytes: int = 20_000) -> bool:
        if not image_url:
            return False
        parsed = urlparse(image_url)
        if parsed.scheme not in {"http", "https"}:
            return False
        try:
            response = self.session.get(image_url, timeout=10, stream=True)
            if response.status_code != 200:
                return False
            content_type = response.headers.get("Content-Type", "").lower()
            if not content_type.startswith("image/"):
                return False
            content_length = response.headers.get("Content-Length")
            if content_length and content_length.isdigit() and int(content_length) < min_bytes:
                return False
            return True
        except Exception:
            return False

    def search_wikimedia(self, query: str) -> Optional[str]:
        """Search Wikimedia Commons API for image URL"""
        api_url = "https://commons.wikimedia.org/w/api.php"
        search_data = self._request_json(
            api_url,
            {
                "action": "query",
                "list": "search",
                "srsearch": f'filetype:bitmap "{query}"',
                "srlimit": 5,
                "format": "json",
            },
        )
        if not search_data:
            return None

        for item in search_data.get("query", {}).get("search", []):
            title = item.get("title", "")
            if not title:
                continue
            if not title.startswith("File:"):
                title = f"File:{title}"
            image_data = self._request_json(
                api_url,
                {
                    "action": "query",
                    "titles": title,
                    "prop": "imageinfo",
                    "iiprop": "url|size",
                    "format": "json",
                },
            )
            if not image_data:
                continue
            for page in image_data.get("query", {}).get("pages", {}).values():
                imageinfo = page.get("imageinfo", [])
                if not imageinfo:
                    continue
                candidate = imageinfo[0].get("url", "")
                if self._validate_image_url(candidate):
                    return candidate
        return None

    def search_serpapi(self, query: str) -> List[str]:
        if not self.serpapi_key:
            return []
        data = self._request_json(
            "https://serpapi.com/search.json",
            {
                "engine": "google_images",
                "q": query,
                "hl": "ru",
                "api_key": self.serpapi_key,
            },
        )
        if not data:
            return []
        urls = []
        for item in data.get("images_results", []):
            original = item.get("original", "")
            if original and self._validate_image_url(original):
                urls.append(original)
            if len(urls) >= 3:
                break
        return urls

    def search_google_cse(self, query: str) -> List[str]:
        if not self.google_cse_key or not self.google_cse_cx:
            return []
        data = self._request_json(
            "https://www.googleapis.com/customsearch/v1",
            {
                "key": self.google_cse_key,
                "cx": self.google_cse_cx,
                "q": query,
                "searchType": "image",
                "num": 5,
                "safe": "active",
            },
        )
        if not data:
            return []
        urls = []
        for item in data.get("items", []):
            link = item.get("link", "")
            if link and self._validate_image_url(link):
                urls.append(link)
            if len(urls) >= 3:
                break
        return urls

    def find_best_photo(self, attraction_name: str, district: str, settlement: str = "") -> Tuple[str, str, float]:
        """Find best photo via API-first providers with confidence"""
        # Build search queries
        queries = []
        
        # Most specific query first
        if settlement:
            queries.append(f"{attraction_name} {settlement} {district} Татарстан")
        
        queries.append(f"{attraction_name} {district} район Татарстан")
        queries.append(f"{attraction_name} Татарстан")
        
        for query in queries:
            print(f"    Searching: {query[:60]}...")
            
            # API-first #1: Wikimedia
            photo_url = self.search_wikimedia(query)
            if photo_url:
                print(f"    ✓ Found on Wikimedia")
                return photo_url, "wikimedia", 0.90

            # API-first #2: SerpAPI (optional)
            serpapi_urls = self.search_serpapi(query)
            if serpapi_urls:
                print(f"    ✓ Found via SerpAPI")
                return serpapi_urls[0], "serpapi", 0.75

            # API-first #3: Google CSE (optional)
            cse_urls = self.search_google_cse(query)
            if cse_urls:
                print(f"    ✓ Found via Google CSE")
                return cse_urls[0], "google_cse", 0.70
            
            time.sleep(1)  # Be polite
        
        print(f"    ✗ No photo found")
        return "", "", 0.0

class OpenRouterClient:
    """Client for OpenRouter AI API"""
    
    def __init__(self, api_key: str, model: str = "meta-llama/llama-3.1-70b-instruct"):
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self.model = model
        self.request_count = 0
        self.token_usage = 0
        self.prompt_version = "v2-structured-extraction"
    
    def chat_completion(self, messages: List[Dict], max_tokens: int = 2000, temperature: float = 0.3) -> str:
        """Send chat completion request"""
        try:
            self.request_count += 1
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                max_tokens=max_tokens,
                temperature=temperature,
            )
            
            # Track token usage
            if hasattr(response, 'usage') and response.usage:
                self.token_usage += response.usage.total_tokens
            
            return response.choices[0].message.content
            
        except Exception as e:
            print(f"  OpenRouter error: {e}")
            time.sleep(2)
            return ""
    
    def extract_attractions_from_text(self, text: str, tables: List[Dict], 
                                     district: str, filename: str) -> List[Dict]:
        """
        Use AI to extract attractions from document text and tables
        """
        # Prepare tables text
        tables_text = ""
        for table in tables:
            tables_text += "\nТАБЛИЦА:\n"
            for row in table.get('rows', []):
                tables_text += " | ".join(str(cell) for cell in row) + "\n"
        
        prompt = f"""
Вы эксперт по извлечению информации о достопримечательностях из документов о районах Татарстана.

РАЙОН: {district}
ФАЙЛ: {filename}

ТЕКСТ ДОКУМЕНТА:
{text}

{tables_text}

ЗАДАЧА:
Извлеките ВСЕ достопримечательности, памятники, музеи, храмы, мечети, природные объекты из этого документа.

Для КАЖДОГО объекта укажите:
1. name - точное название достопримечательности
2. settlement - населенный пункт (село, город, поселок), где находится (если указан)
3. brief_description - краткое описание 1-2 предложения ТОЛЬКО из документа
4. type - тип объекта (музей, памятник, храм, мечеть, природный объект, усадьба и т.д.)
5. confidence - число от 0 до 1, насколько уверены в корректности извлечения
6. source_fragment - короткая цитата (до 200 символов) из документа, подтверждающая объект

Верните ТОЛЬКО валидный JSON массив в формате:
[
  {{
    "name": "Название достопримечательности",
    "settlement": "Название населенного пункта",
    "brief_description": "Краткое описание",
    "type": "категория",
    "confidence": 0.82,
    "source_fragment": "Фрагмент текста документа"
  }}
]

Если достопримечательностей не найдено, верните пустой массив [].
Если документ содержит только список без деталей, извлеките что можете.
Если в документе одна большая достопримечательность - извлеките ее.
"""
        
        messages = [
            {"role": "system", "content": "Вы полезный помощник, который извлекает структурированные данные из документов. Всегда отвечайте валидным JSON."},
            {"role": "user", "content": prompt}
        ]
        
        response = self.chat_completion(messages, max_tokens=3500)
        
        # Parse JSON from response
        attractions = self._parse_json_response(response)
        
        if not attractions and text.strip():
            # Fallback: create single attraction from document
            attractions = [{
                "name": filename.replace('.docx', '').replace('.doc', ''),
                "settlement": "",
                "brief_description": text[:500] if text else "Информация о достопримечательности",
                "type": "не определено",
                "confidence": 0.35,
                "source_fragment": text[:200] if text else ""
            }]
        
        return attractions if attractions else []
    
    def _parse_json_response(self, response: str) -> List[Dict]:
        """Parse JSON from AI response"""
        try:
            # Find JSON array in response
            json_match = re.search(r'\[\s*\{.*\}\s*\]', response, re.DOTALL)
            if json_match:
                attractions = json.loads(json_match.group())
                return attractions
            else:
                # Try to parse entire response
                attractions = json.loads(response)
                return attractions
        except json.JSONDecodeError as e:
            print(f"  JSON parse error: {e}")
            # Try to find any JSON-like structure
            try:
                # Look for objects
                obj_pattern = r'\{[^{}]*"name"[^{}]*\}'
                matches = re.findall(obj_pattern, response, re.DOTALL)
                if matches:
                    # Try to parse each match
                    result = []
                    for match in matches:
                        try:
                            obj = json.loads(match)
                            result.append(obj)
                        except:
                            pass
                    return result
            except:
                pass
            return []
    
    def enrich_attraction_description(self, attraction: Dict, district: str) -> Dict:
        """
        Enrich attraction with detailed HTML description using AI
        """
        name = attraction.get('name', '')
        settlement = attraction.get('settlement', '')
        brief = attraction.get('brief_description', '')
        attr_type = attraction.get('type', '')
        
        location_info = f"{settlement}, {district}" if settlement else district
        
        prompt = f"""
Вы эксперт-краевед по Республике Татарстан.

ДОСТОПРИМЕЧАТЕЛЬНОСТЬ: {name}
РАСПОЛОЖЕНИЕ: {location_info}, Республика Татарстан
ТИП: {attr_type}
КРАТКАЯ ИНФОРМАЦИЯ: {brief}

ЗАДАЧА:
1. Создайте подробное описание этой достопримечательности (3-4 абзаца в формате HTML)
2. Включите историческую справку, архитектурные особенности, культурную значимость
3. Если точной информации нет, дайте общий контекст о подобных объектах в Татарстане
4. Пишите на русском языке
5. Используйте правильную HTML разметку с тегами <p>
6. Добавьте интересные факты если возможно

Верните JSON:
{{
  "description_html": "<p>Первый абзац...</p><p>Второй абзац...</p><p>Третий абзац...</p>",
  "interesting_facts": ["Интересный факт 1", "Интересный факт 2"],
  "visiting_info": "Информация о посещении (если есть)",
  "historical_period": "Исторический период/век"
}}
"""
        
        messages = [
            {"role": "system", "content": "Вы знающий краевед и гид. Предоставляйте точную, увлекательную информацию."},
            {"role": "user", "content": prompt}
        ]
        
        response = self.chat_completion(messages, max_tokens=2500)
        
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                enrichment = json.loads(json_match.group())
                return enrichment
        except:
            pass
        
        # Fallback: generate simple HTML from brief description
        paragraphs = []
        if brief:
            sentences = re.split(r'[.!?]+', brief)
            para1 = ". ".join(s.strip() for s in sentences[:3] if s.strip())
            para2 = ". ".join(s.strip() for s in sentences[3:6] if s.strip())
            para3 = f"Данная достопримечательность расположена в {location_info} районе Республики Татарстан и представляет культурно-историческую ценность для региона."
            
            if para1:
                paragraphs.append(f"<p>{para1}.</p>")
            if para2:
                paragraphs.append(f"<p>{para2}.</p>")
            paragraphs.append(f"<p>{para3}</p>")
        else:
            paragraphs = [
                f"<p>{name} - достопримечательность {location_info} района Республики Татарстан.</p>",
                f"<p>Объект представляет интерес для туристов и краеведов, изучающих историю и культуру региона.</p>",
                f"<p>Рекомендуется к посещению всем, кто интересуется историческим наследием Татарстана.</p>"
            ]
        
        return {
            "description_html": "".join(paragraphs),
            "interesting_facts": [],
            "visiting_info": "",
            "historical_period": ""
        }

    def regenerate_description_strict(self, attraction: Dict, district: str, min_paragraphs: int = 2) -> Dict:
        """Regenerate HTML description with strict quality constraints"""
        name = attraction.get('name', '')
        settlement = attraction.get('settlement', '')
        brief = attraction.get('brief_description', '')
        attr_type = attraction.get('type', '')
        location_info = f"{settlement}, {district}" if settlement else district

        prompt = f"""
Сформируйте качественное и нейтральное описание достопримечательности.

Название: {name}
Локация: {location_info}, Республика Татарстан
Тип: {attr_type}
Известные данные: {brief}

Требования:
- Только русский язык
- Минимум {min_paragraphs} абзаца в HTML <p>...</p>
- Без вымышленных деталей, только осторожные формулировки при нехватке данных
- Без английских слов внутри русских предложений

Верните JSON:
{{
  "description_html": "<p>...</p><p>...</p>",
  "interesting_facts": [],
  "visiting_info": "",
  "historical_period": ""
}}
"""
        messages = [
            {"role": "system", "content": "Верни строго JSON. Следуй требованиям к качеству текста."},
            {"role": "user", "content": prompt}
        ]
        response = self.chat_completion(messages, max_tokens=1500, temperature=0.2)
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
        except Exception:
            pass
        return self.enrich_attraction_description(attraction, district)

class AttractionProcessor:
    """Main processor for attractions"""
    
    CACHE_VERSION = "2.0"

    def __init__(
        self,
        openrouter_api_key: str,
        districts_json: str = None,
        model: str = "meta-llama/llama-3.1-70b-instruct",
        skip_photos: bool = False,
        min_description_paragraphs: int = 2,
        retry_count: int = 2,
        resume: bool = False,
        max_files: int = 0,
        only_district: str = "",
    ):
        self.parser = DOCXParser()
        self.ai_client = OpenRouterClient(openrouter_api_key, model=model)
        self.image_searcher = ImageSearcher()
        self.normalizer = DistrictNormalizer(districts_json)
        self.processed_count = 0
        self.processed_files_count = 0
        self.success_files_count = 0
        self.cache_file = "processing_cache.json"
        self.state_file = "processing_state.json"
        self.error_file = "processing_errors.json"
        self.skip_photos = skip_photos
        self.min_description_paragraphs = min_description_paragraphs
        self.retry_count = retry_count
        self.max_files = max_files
        self.only_district = only_district.lower().strip()
        self.resume = resume
        self.errors: List[Dict] = []
        self.quality_stats = {
            "with_photo": 0,
            "with_min_paragraphs": 0,
            "high_confidence": 0,
        }
        self.config_signature = hashlib.md5(
            json.dumps(
                {
                    "model": model,
                    "prompt_version": self.ai_client.prompt_version,
                    "min_description_paragraphs": min_description_paragraphs,
                    "skip_photos": skip_photos,
                    "retry_count": retry_count,
                },
                sort_keys=True,
                ensure_ascii=False,
            ).encode("utf-8")
        ).hexdigest()
        self.cache = self._load_cache()
        self.state = self._load_state()
    
    def _load_cache(self) -> Dict:
        """Load processing cache"""
        try:
            if os.path.exists(self.cache_file):
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                if isinstance(data, dict) and data.get("_meta"):
                    return data
                return {"_meta": {}, "items": data if isinstance(data, dict) else {}}
        except:
            pass
        return {"_meta": {}, "items": {}}
    
    def _save_cache(self):
        """Save processing cache"""
        self.cache["_meta"] = {
            "version": self.CACHE_VERSION,
            "config_signature": self.config_signature,
            "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
        except:
            pass

    def _load_state(self) -> Dict:
        if not self.resume or not os.path.exists(self.state_file):
            return {"processed_files": []}
        try:
            with open(self.state_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, dict) and "processed_files" in data:
                    return data
        except Exception:
            pass
        return {"processed_files": []}

    def _save_state(self):
        data = {
            "processed_files": self.state.get("processed_files", []),
            "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        try:
            with open(self.state_file, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _save_errors(self):
        try:
            with open(self.error_file, "w", encoding="utf-8") as f:
                json.dump(self.errors, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _log_error(self, file_path: str, error_type: str, message: str):
        self.errors.append(
            {
                "file": file_path,
                "error_type": error_type,
                "message": message,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
        )
    
    def _get_file_hash(self, file_path: str) -> str:
        """Get hash of file content"""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except:
            return ""

    def _count_html_paragraphs(self, html: str) -> int:
        return len(re.findall(r"<p>.*?</p>", html, flags=re.DOTALL))

    def _has_language_artifacts(self, html: str) -> bool:
        latin_words = re.findall(r"\b[A-Za-z]{4,}\b", html)
        return len(latin_words) > 0

    def _normalize_settlement(self, district: str, settlement: str) -> Tuple[str, str]:
        raw_settlement = settlement or ""
        if not raw_settlement:
            return "", ""
        settlements = self.normalizer.settlement_map.get(district, [])
        raw_norm = raw_settlement.lower().strip()
        for known in settlements:
            if known.lower() == raw_norm:
                return known, raw_settlement
            if known.lower() in raw_norm or raw_norm in known.lower():
                return known, raw_settlement
        return raw_settlement, raw_settlement

    def _validate_extracted_attraction(self, attr_data: Dict) -> Optional[Dict]:
        name = str(attr_data.get("name", "")).strip()
        if len(name) < 3:
            return None
        settlement = str(attr_data.get("settlement", "")).strip()
        brief = str(attr_data.get("brief_description", "")).strip()
        attr_type = str(attr_data.get("type", "")).strip() or "не определено"
        source_fragment = str(attr_data.get("source_fragment", "")).strip()[:200]
        confidence_raw = attr_data.get("confidence", 0.5)
        try:
            confidence = float(confidence_raw)
        except (TypeError, ValueError):
            confidence = 0.5
        confidence = max(0.0, min(1.0, confidence))
        return {
            "name": name,
            "settlement": settlement,
            "brief_description": brief,
            "type": attr_type,
            "confidence": confidence,
            "source_fragment": source_fragment,
        }
    
    def process_file(self, file_path: str, district_folder: str) -> List[Attraction]:
        """Process single DOCX file"""
        print(f"\n  Файл: {os.path.basename(file_path)}")
        self.processed_files_count += 1
        
        # Check cache
        file_hash = self._get_file_hash(file_path)
        cache_key = f"{district_folder}:{os.path.basename(file_path)}"
        cache_meta = self.cache.get("_meta", {})
        cache_items = self.cache.get("items", {})
        
        if (
            cache_meta.get("version") == self.CACHE_VERSION
            and cache_meta.get("config_signature") == self.config_signature
            and cache_key in cache_items
            and cache_items[cache_key].get('hash') == file_hash
        ):
            print(f"    Используем кэш ({len(cache_items[cache_key].get('attractions', []))} объектов)")
            cached_data = cache_items[cache_key].get('attractions', [])
            self.success_files_count += 1
            return [Attraction(**attr) for attr in cached_data]
        
        # Normalize district name
        district = self.normalizer.normalize(district_folder)
        
        # Extract content from DOCX/DOC
        text, tables, title, parse_error = self.parser.extract_content(file_path)
        
        if not text and not tables:
            print("    Нет содержимого")
            if parse_error:
                self._log_error(file_path, "parse_error", parse_error)
            return []
        
        print(f"    Извлечено текста: {len(text)} символов, таблиц: {len(tables)}")
        
        # Extract attractions using AI
        attractions_data = self.ai_client.extract_attractions_from_text(
            text, tables, district, 
            title or os.path.basename(file_path)
        )
        
        if not attractions_data:
            print("    Достопримечательности не извлечены")
            self._log_error(file_path, "extract_error", "AI не извлек объекты из документа")
            return []
        
        cleaned_attractions_data = []
        for item in attractions_data:
            cleaned = self._validate_extracted_attraction(item)
            if cleaned:
                cleaned_attractions_data.append(cleaned)

        if not cleaned_attractions_data:
            self._log_error(file_path, "validation_error", "После валидации не осталось валидных объектов")
            return []

        print(f"    Найдено объектов: {len(cleaned_attractions_data)}")
        
        # Enrich each attraction
        attractions = []
        for i, attr_data in enumerate(cleaned_attractions_data):
            print(f"    Обработка {i+1}/{len(cleaned_attractions_data)}: {attr_data.get('name', 'Unknown')[:50]}")
            
            # Find settlement in text
            settlement = attr_data.get('settlement', '')
            if not settlement:
                settlement = self.normalizer.find_settlement(district, text)
            settlement, raw_settlement = self._normalize_settlement(district, settlement)
            
            # Enrich with AI
            enrichment = {}
            for attempt in range(self.retry_count + 1):
                enrichment = self.ai_client.enrich_attraction_description(attr_data, district)
                paragraph_count = self._count_html_paragraphs(enrichment.get("description_html", ""))
                has_artifacts = self._has_language_artifacts(enrichment.get("description_html", ""))
                if paragraph_count >= self.min_description_paragraphs and not has_artifacts:
                    break
                if attempt < self.retry_count:
                    enrichment = self.ai_client.regenerate_description_strict(
                        attr_data, district, self.min_description_paragraphs
                    )
            
            # Find photo
            photo_url = ""
            photo_source = ""
            photo_confidence = 0.0
            if not self.skip_photos:
                photo_url, photo_source, photo_confidence = self.image_searcher.find_best_photo(
                    attr_data.get('name', ''),
                    district,
                    settlement
                )
            
            # Create attraction object
            description_html = enrichment.get('description_html', '')
            paragraph_count = self._count_html_paragraphs(description_html)
            if paragraph_count >= self.min_description_paragraphs:
                self.quality_stats["with_min_paragraphs"] += 1
            if photo_url:
                self.quality_stats["with_photo"] += 1
            if attr_data.get("confidence", 0) >= 0.7:
                self.quality_stats["high_confidence"] += 1

            attraction = Attraction(
                name=attr_data.get('name', ''),
                district=district,
                settlement=settlement,
                description_html=description_html,
                photo_url=photo_url,
                photo_source=photo_source,
                photo_confidence=photo_confidence,
                additional_info={
                    'type': attr_data.get('type', ''),
                    'brief_description': attr_data.get('brief_description', ''),
                    'source_fragment': attr_data.get('source_fragment', ''),
                    'confidence': attr_data.get('confidence', 0.5),
                    'raw_settlement': raw_settlement,
                    'interesting_facts': enrichment.get('interesting_facts', []),
                    'visiting_info': enrichment.get('visiting_info', ''),
                    'historical_period': enrichment.get('historical_period', '')
                }
            )
            
            attractions.append(attraction)
            self.processed_count += 1
            
            # Rate limiting
            time.sleep(0.5)
        
        # Save to cache
        if attractions:
            self.cache.setdefault("items", {})[cache_key] = {
                'hash': file_hash,
                'attractions': [attr.to_dict() for attr in attractions]
            }
            self._save_cache()
            self.success_files_count += 1
        
        return attractions
    
    def process_directory(self, root_folder: str) -> List[Attraction]:
        """Process all DOCX files in directory structure"""
        all_attractions = []
        root_path = Path(root_folder)
        
        # Get all district folders
        district_folders = [d for d in root_path.iterdir() if d.is_dir()]
        
        print(f"\nНайдено папок районов: {len(district_folders)}")
        print("="*70)
        
        for idx, district_folder in enumerate(district_folders, 1):
            district_name = district_folder.name
            if self.only_district and self.only_district not in district_name.lower():
                continue
            print(f"\n[{idx}/{len(district_folders)}] РАЙОН: {district_name}")
            print("-"*70)
            
            # Get all DOCX files
            docx_files = list(district_folder.glob("*.docx")) + list(district_folder.glob("*.doc"))
            
            if not docx_files:
                print(f"  Нет DOCX файлов")
                continue
            
            print(f"  Файлов: {len(docx_files)}")
            
            # Process each file
            for docx_file in docx_files:
                state_key = str(docx_file.resolve())
                if self.resume and state_key in self.state.get("processed_files", []):
                    print("  Пропуск по resume-state")
                    continue
                if self.max_files and self.processed_files_count >= self.max_files:
                    print("  Достигнут лимит --max-files")
                    return all_attractions
                try:
                    attractions = self.process_file(str(docx_file), district_name)
                    all_attractions.extend(attractions)
                    self.state.setdefault("processed_files", []).append(state_key)
                    self._save_state()
                except Exception as e:
                    print(f"  Ошибка обработки файла: {e}")
                    self._log_error(str(docx_file), "runtime_error", str(e))
                    self._save_errors()
                
                # Save progress every 5 files
                if len(all_attractions) % 5 == 0:
                    self.save_progress(all_attractions, "progress_backup.json")
        
        return all_attractions
    
    def save_progress(self, attractions: List[Attraction], filename: str):
        """Save intermediate results"""
        data = [attr.to_dict() for attr in attractions]
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"\n  Прогресс сохранен в {filename} ({len(attractions)} объектов)")
    
    def export_json(self, attractions: List[Attraction], output_file: str):
        """Export final results to JSON"""
        # Group by district
        districts = {}
        for attr in attractions:
            if attr.district not in districts:
                districts[attr.district] = []
            districts[attr.district].append(attr.to_dict())

        total = len(attractions)
        quality = {
            "attractions_with_photo_percent": round((self.quality_stats["with_photo"] / total * 100), 2) if total else 0,
            "descriptions_with_min_paragraphs_percent": round((self.quality_stats["with_min_paragraphs"] / total * 100), 2) if total else 0,
            "high_confidence_percent": round((self.quality_stats["high_confidence"] / total * 100), 2) if total else 0,
            "processed_files": self.processed_files_count,
            "successful_files": self.success_files_count,
            "file_success_percent": round((self.success_files_count / self.processed_files_count * 100), 2) if self.processed_files_count else 0,
            "errors_count": len(self.errors),
        }
        
        data = {
            "metadata": {
                "total_attractions": len(attractions),
                "total_districts": len(districts),
                "processed_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "ai_model": self.ai_client.model,
                "ai_requests": self.ai_client.request_count,
                "estimated_tokens": self.ai_client.token_usage,
                "cache_version": self.CACHE_VERSION,
                "prompt_version": self.ai_client.prompt_version,
                "districts": list(districts.keys())
            },
            "quality": quality,
            "attractions": [attr.to_dict() for attr in attractions],
            "by_district": districts
        }
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        quality_file = f"{Path(output_file).stem}_quality_report.json"
        with open(quality_file, "w", encoding="utf-8") as f:
            json.dump(quality, f, ensure_ascii=False, indent=2)
        self._save_errors()
        
        print(f"\n{'='*70}")
        print(f"ЭКСПОРТ ЗАВЕРШЕН")
        print(f"{'='*70}")
        print(f"Всего достопримечательностей: {len(attractions)}")
        print(f"Районов: {len(districts)}")
        print(f"Файл: {output_file}")
        print(f"Quality report: {quality_file}")
        print(f"С фото: {quality['attractions_with_photo_percent']}%")
        print(f"Описание >= {self.min_description_paragraphs} абз.: {quality['descriptions_with_min_paragraphs_percent']}%")
        print(f"Высокая уверенность: {quality['high_confidence_percent']}%")
        print(f"{'='*70}")

def main():
    """Main entry point"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Парсинг достопримечательностей из DOCX файлов')
    parser.add_argument('--input', '-i', required=True, help='Папка с подпапками районов')
    parser.add_argument('--output', '-o', default='attractions.json', help='Выходной JSON файл')
    parser.add_argument('--api-key', required=True, help='OpenRouter API ключ')
    parser.add_argument('--model', default='meta-llama/llama-3.1-70b-instruct', 
                       help='Модель OpenRouter')
    parser.add_argument('--districts', '-d', help='JSON файл со списком районов и населенных пунктов')
    parser.add_argument('--max-files', type=int, default=0, help='Ограничить количество обрабатываемых файлов (0 - без лимита)')
    parser.add_argument('--only-district', default='', help='Обрабатывать только районы, содержащие эту строку')
    parser.add_argument('--skip-photos', action='store_true', help='Пропустить поиск фотографий')
    parser.add_argument('--min-description-paragraphs', type=int, default=2, help='Минимум абзацев в HTML-описании')
    parser.add_argument('--retry-count', type=int, default=2, help='Количество повторов для регенерации описания')
    parser.add_argument('--resume', action='store_true', help='Продолжить обработку по processing_state.json')
    
    args = parser.parse_args()
    
    # Load districts data if provided
    districts_json = None
    if args.districts and os.path.exists(args.districts):
        with open(args.districts, 'r', encoding='utf-8') as f:
            districts_json = f.read()
        print(f"Загружен список районов из {args.districts}")
    
    # Initialize processor
    print(f"\n{'='*70}")
    print(f"ПАРСЕР ДОСТОПРИМЕЧАТЕЛЬНОСТЕЙ ТАТАРСТАНА")
    print(f"{'='*70}")
    print(f"Входная папка: {args.input}")
    print(f"Выходной файл: {args.output}")
    print(f"Модель AI: {args.model}")
    print(f"Поиск фото: {'выкл' if args.skip_photos else 'вкл'}")
    print(f"Resume: {'да' if args.resume else 'нет'}")
    print(f"{'='*70}\n")
    
    processor = AttractionProcessor(
        args.api_key,
        districts_json,
        model=args.model,
        skip_photos=args.skip_photos,
        min_description_paragraphs=max(1, args.min_description_paragraphs),
        retry_count=max(0, args.retry_count),
        resume=args.resume,
        max_files=max(0, args.max_files),
        only_district=args.only_district,
    )
    
    # Process all files
    attractions = processor.process_directory(args.input)
    
    if not attractions:
        print("\n⚠️  Достопримечательности не найдены!")
        return
    
    # Export results
    processor.export_json(attractions, args.output)
    
    print(f"\nВсего обработано: {processor.processed_count} достопримечательностей")
    print(f"Запросов к AI: {processor.ai_client.request_count}")
    print(f"Использовано токенов: ~{processor.ai_client.token_usage}")

if __name__ == "__main__":
    main()
