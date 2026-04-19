# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import subprocess
import tempfile
import time
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

from docx import Document


def _default_subprocess_run(cmd: List[str], **kwargs: Any) -> Any:
    return subprocess.run(cmd, **kwargs)


class DOCXParser:
    """Parser for DOCX files with tables support"""

    def __init__(
        self,
        subprocess_run: Optional[Callable[..., Any]] = None,
    ) -> None:
        self._subprocess_run = subprocess_run or _default_subprocess_run

    def _convert_doc_to_docx(self, file_path: str) -> Optional[str]:
        """Convert legacy .doc file to .docx using LibreOffice"""
        with tempfile.TemporaryDirectory(prefix="doc_convert_") as temp_dir:
            command = [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                temp_dir,
                file_path,
            ]
            result = self._subprocess_run(command, capture_output=True, text=True, timeout=90)
            if result.returncode != 0:
                return None
            converted = Path(temp_dir) / f"{Path(file_path).stem}.docx"
            if not converted.exists():
                return None
            final_path = Path(tempfile.gettempdir()) / f"{Path(file_path).stem}_{int(time.time() * 1000)}.docx"
            final_path.write_bytes(converted.read_bytes())
            return str(final_path)

    def extract_content(self, file_path: str) -> Tuple[str, List[Dict[str, Any]], str, Optional[str]]:
        """
        Extract text, tables and metadata from DOCX
        Returns: (full_text, tables_data, title, error_message)
        """
        source_file = file_path
        temp_converted_file: Optional[str] = None
        try:
            extension = Path(file_path).suffix.lower()
            if extension == ".doc":
                temp_converted_file = self._convert_doc_to_docx(file_path)
                if not temp_converted_file:
                    return "", [], "", "Не удалось конвертировать .doc в .docx через soffice"
                source_file = temp_converted_file

            doc = Document(source_file)
            paragraphs: List[str] = []
            tables_data: List[Dict[str, Any]] = []
            title = ""

            # Extract title from first paragraph if it looks like a title
            if doc.paragraphs:
                first_para = doc.paragraphs[0].text.strip()
                if len(first_para) < 200 and first_para:
                    title = first_para

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data: Dict[str, Any] = {"index": table_idx, "rows": []}

                for row_idx, row in enumerate(table.rows):
                    row_data: List[str] = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    if row_data:
                        table_data["rows"].append(row_data)
                        # Keep table text in full-text fallback for LLM extraction
                        paragraphs.append(" | ".join(row_data))

                if table_data["rows"]:
                    tables_data.append(table_data)

            full_text = "\n".join(paragraphs)
            return full_text, tables_data, title, None

        except Exception as e:
            return "", [], "", str(e)
        finally:
            if temp_converted_file and os.path.exists(temp_converted_file):
                try:
                    os.remove(temp_converted_file)
                except OSError:
                    pass
